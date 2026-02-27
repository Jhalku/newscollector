import os
import sys
import io
import time
import requests
from datetime import datetime, timedelta
from dotenv import load_dotenv
from GoogleNews import GoogleNews
from docx import Document
from docx.shared import Pt
import argparse

# set stdout to utf-8 to avoid charmap errors
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

load_dotenv()

TELEGRAM_BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
TELEGRAM_CHAT_ID = os.getenv('TELEGRAM_CHAT_ID')

DEFAULT_KEYWORDS = os.getenv('NEWS_KEYWORDS', 'AI India').split(',')
DEFAULT_SITES = os.getenv('NEWS_SITES', 'hindustantimes.com,bbc.com').split(',')

# Clean up lists
DEFAULT_KEYWORDS = [k.strip() for k in DEFAULT_KEYWORDS if k.strip()]
DEFAULT_SITES = [s.strip() for s in DEFAULT_SITES if s.strip()]

def fetch_news(keywords, sites, dry_run=False):
    import random
    import xml.etree.ElementTree as ET
    from urllib.parse import quote
    
    end_date = datetime.now()
    # Search for the last 48 hours to get recent news
    start_date = end_date - timedelta(hours=48)
    
    start_str = start_date.strftime('%m/%d/%Y')
    end_str = end_date.strftime('%m/%d/%Y')
    
    print(f"Searching from {start_str} to {end_str} via RSS...")
    
    all_results = []
    
    # RSS feeds are much less likely to 429
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    }
    
    for keyword in keywords:
        for site in sites:
            # query parameter: "keyword" site:site.com when:2d
            query = f'"{keyword}" site:{site} when:2d'
            print(f"Checking RSS for query: {query}")
            
            try:
                # Direct Google News RSS search URL
                url = f"https://news.google.com/rss/search?q={quote(query)}&hl=en-US&gl=US&ceid=US:en"
                
                response = requests.get(url, headers=headers, timeout=10)
                
                if response.status_code == 429:
                    print(f"Still hitting rate limit for {query} on RSS ! Pausing 15s...")
                    time.sleep(15)
                    continue
                    
                response.raise_for_status()
                
                # Parse the XML response
                root = ET.fromstring(response.text)
                channel = root.find('channel')
                
                if not channel:
                    continue
                    
                unique_articles = []
                seen_links = set()
                
                items = channel.findall('item')
                for item in items:
                    if len(unique_articles) >= 5:
                        break
                        
                    title = item.findtext('title', default='News Article')
                    link = item.findtext('link', default='')
                    published = item.findtext('pubDate', default='Recent')
                    # RSS description contains HTML, we just want a snippet
                    desc = item.findtext('description', default='')
                    
                    if not link or link in seen_links:
                        continue
                        
                    # Title typically comes as "Article Title - Source Name"
                    # Filter by keyword presence to ensure relevance
                    if keyword.lower() in title.lower():
                        seen_links.add(link)
                        unique_articles.append({
                            'title': title,
                            'link': link,
                            'desc': "Retrieved via Google News RSS.", # The RSS desc is messy HTML, keep it clean
                            'published': published,
                            'site': site,
                            'keyword': keyword
                        })
                
                all_results.extend(unique_articles)
                
            except Exception as e:
                print(f"Error fetching RSS for {query}: {e}")
                
            # Random delay between combinations
            sleep_time = random.uniform(1.0, 3.0)
            time.sleep(sleep_time)
            
    return all_results

def generate_word_doc(articles, filepath):
    doc = Document()
    doc.add_heading(f"News Report - {datetime.now().strftime('%Y-%m-%d')}", 0)
    
    if not articles:
        doc.add_paragraph("No specific news found for the given keywords and sites in the recent timeframe.")
    else:
        # Group by keyword
        grouped = {}
        for a in articles:
            k = a['keyword']
            if k not in grouped:
                grouped[k] = []
            grouped[k].append(a)
            
        for keyword, group_articles in grouped.items():
            doc.add_heading(f"Keyword: {keyword}", level=1)
            for i, article in enumerate(group_articles, 1):
                p_title = doc.add_paragraph(style='List Number')
                runner = p_title.add_run(f"[{article['site']}] {article['title']}")
                runner.bold = True
                
                p_desc = doc.add_paragraph(f"{article['desc']} ({article['published']})")
                p_desc.style.font.size = Pt(10)
                
                p_link = doc.add_paragraph()
                # Create rudimentary clickable link format (Word parses these if URLs)
                p_link.add_run(article['link'])
                
    doc.save(filepath)
    print(f"Saved Word document to {filepath}")

def send_telegram_message(articles, filepath):
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        print("Telegram details not configured. Skipping message.")
        return
        
    date_str = datetime.now().strftime('%Y-%m-%d')
    message = f"📰 <b>Daily News Report ({date_str})</b>\n\n"
    
    if not articles:
        message += "No new articles found today."
    else:
        # Just send top 10 overall in text to avoid message too long errors
        top_articles = articles[:10]
        for a in top_articles:
            title = a['title']
            # Escape HTML characters for Telegram
            title = title.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            message += f"▪️ <a href='{a['link']}'>{title}</a> ({a['site']})\n"
            
        if len(articles) > 10:
            message += f"\n<i>...and {len(articles) - 10} more. See the attached Word document!</i>"
    
    url_msg = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    payload_msg = {
        'chat_id': TELEGRAM_CHAT_ID,
        'text': message,
        'parse_mode': 'HTML',
        'disable_web_page_preview': True
    }
    
    print("Sending Telegram summary...")
    try:
        requests.post(url_msg, data=payload_msg)
    except Exception as e:
        print(f"Failed to send Telegram message: {e}")
        
    # Send document
    url_doc = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument"
    print("Sending Telegram document...")
    try:
        with open(filepath, 'rb') as doc_file:
            files = {'document': doc_file}
            data = {'chat_id': TELEGRAM_CHAT_ID}
            requests.post(url_doc, data=data, files=files)
    except Exception as e:
        print(f"Failed to send Telegram document: {e}")
    
    print("Telegram delivery complete.")

def main():
    parser = argparse.ArgumentParser(description="News Tracker Tool")
    parser.add_argument("--test", action="store_true", help="Run a dry run without sending telegram messages")
    args = parser.parse_args()
    
    if not DEFAULT_KEYWORDS or not DEFAULT_SITES:
        print("Please configure NEWS_KEYWORDS and NEWS_SITES in .env file.")
        return
        
    print(f"Starting news tracker for keywords: {DEFAULT_KEYWORDS}")
    print(f"Monitoring sites: {DEFAULT_SITES}")
    
    articles = fetch_news(DEFAULT_KEYWORDS, DEFAULT_SITES, dry_run=args.test)
    
    print(f"\nFound {len(articles)} total unique articles.")
    
    # Ensure output directory exists
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
    os.makedirs(output_dir, exist_ok=True)
    
    date_str = datetime.now().strftime('%Y%m%d')
    doc_path = os.path.join(output_dir, f"News_Report_{date_str}.docx")
    
    generate_word_doc(articles, doc_path)
    
    if not args.test:
        send_telegram_message(articles, doc_path)
    else:
        print("Test mode enabled. Skipping Telegram delivery.")
        if articles:
            print("\nSample top link:")
            print(f"Title: {articles[0]['title']}")
            print(f"Link: {articles[0]['link']}")

if __name__ == "__main__":
    main()
